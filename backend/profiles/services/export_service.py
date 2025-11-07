"""
Export service for generating PDF and DOCX files from CV data.
"""
import os
import tempfile
from pathlib import Path
from typing import Dict, Any, Tuple
from django.template.loader import render_to_string
from django.conf import settings
from playwright.sync_api import sync_playwright
import logging

logger = logging.getLogger(__name__)


class CVExportService:
    """Service for exporting CV to various formats."""
    
    TEMPLATE_MAP = {
        'clean': 'cv/clean.html',
        'two-column': 'cv/two-column.html',
    }
    
    @classmethod
    def _prepare_context(cls, cv) -> Dict[str, Any]:
        """Prepare context data for template rendering."""
        sections = cv.sections or {}
        
        # Extract data from sections
        personal = sections.get('personal', {})
        summary = sections.get('summary', '')
        education = sections.get('education', [])
        experience = sections.get('experience', [])
        skills = sections.get('skills', [])
        projects = sections.get('projects', [])
        links = sections.get('links', {})
        
        return {
            'cv': cv,
            'personal': personal,
            'summary': summary,
            'education': education,
            'experience': experience,
            'skills': skills,
            'projects': projects,
            'links': links,
        }
    
    @classmethod
    def generate_html(cls, cv) -> str:
        """Generate HTML from CV data."""
        template_key = cv.template_key or 'clean'
        template_name = cls.TEMPLATE_MAP.get(template_key, cls.TEMPLATE_MAP['clean'])
        
        context = cls._prepare_context(cv)
        html_content = render_to_string(template_name, context)
        
        return html_content
    
    @classmethod
    def generate_pdf(cls, cv) -> Tuple[bytes, str]:
        """
        Generate PDF from CV using Playwright.
        Returns tuple of (pdf_bytes, filename).
        """
        try:
            html_content = cls.generate_html(cv)
            
            # Create temporary HTML file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_html:
                temp_html.write(html_content)
                temp_html_path = temp_html.name
            
            try:
                # Generate PDF using Playwright
                with sync_playwright() as p:
                    browser = p.chromium.launch(headless=True)
                    page = browser.new_page()
                    
                    # Load HTML file
                    page.goto(f'file://{temp_html_path}')
                    
                    # Wait for page to load
                    page.wait_for_load_state('networkidle')
                    
                    # Generate PDF
                    pdf_bytes = page.pdf(
                        format='A4',
                        print_background=True,
                        margin={
                            'top': '20mm',
                            'right': '20mm',
                            'bottom': '20mm',
                            'left': '20mm',
                        }
                    )
                    
                    browser.close()
                
                # Generate filename
                filename = f"cv-{cv.id}-v{cv.version}.pdf"
                
                logger.info(f"Successfully generated PDF for CV {cv.id}")
                return pdf_bytes, filename
                
            finally:
                # Clean up temporary HTML file
                try:
                    os.unlink(temp_html_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temporary HTML file: {e}")
                    
        except Exception as e:
            logger.error(f"Error generating PDF for CV {cv.id}: {e}", exc_info=True)
            raise
    
    @classmethod
    def generate_docx(cls, cv) -> Tuple[bytes, str]:
        """
        Generate DOCX from CV using python-docx.
        Returns tuple of (docx_bytes, filename).
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Set up margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Get CV data
            cv_data = cv.sections or {}
            personal = cv_data.get('personal', {})
            summary = cv_data.get('summary', '')
            education = cv_data.get('education', [])
            experience = cv_data.get('experience', [])
            skills = cv_data.get('skills', [])
            projects = cv_data.get('projects', [])
            links = cv_data.get('links', {})
            
            # Add header with name
            name = personal.get('name', 'Your Name')
            heading = doc.add_heading(name, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact info
            contact_parts = []
            if personal.get('email'):
                contact_parts.append(personal['email'])
            if personal.get('phone'):
                contact_parts.append(personal['phone'])
            if personal.get('location'):
                contact_parts.append(personal['location'])
            
            if contact_parts:
                contact = doc.add_paragraph(' • '.join(contact_parts))
                contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact.runs[0].font.size = Pt(10)
            
            # Links
            if links:
                link_parts = []
                for key, value in links.items():
                    if value:
                        link_parts.append(value)
                if link_parts:
                    links_para = doc.add_paragraph(' | '.join(link_parts))
                    links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    links_para.runs[0].font.size = Pt(10)
            
            doc.add_paragraph()  # Spacing
            
            # Summary
            if summary:
                doc.add_heading('Summary', level=2)
                doc.add_paragraph(summary)
            
            # Experience
            if experience:
                doc.add_heading('Experience', level=2)
                for exp in experience:
                    # Position and dates
                    position = exp.get('position', 'Position')
                    company = exp.get('company', 'Company')
                    start_date = exp.get('start_date', '')
                    end_date = exp.get('end_date', 'Present')
                    
                    p = doc.add_paragraph()
                    run = p.add_run(position)
                    run.bold = True
                    run.font.size = Pt(11)
                    
                    # Company and dates
                    p = doc.add_paragraph()
                    p.add_run(f"{company} | {start_date} - {end_date}").italic = True
                    
                    # Description
                    if exp.get('description'):
                        doc.add_paragraph(exp['description'])
                    
                    # Achievements
                    if exp.get('achievements'):
                        for achievement in exp['achievements']:
                            p = doc.add_paragraph(achievement, style='List Bullet')
                            p.paragraph_format.left_indent = Inches(0.25)
                    
                    doc.add_paragraph()  # Spacing
            
            # Education
            if education:
                doc.add_heading('Education', level=2)
                for edu in education:
                    degree = edu.get('degree', 'Degree')
                    institution = edu.get('institution', 'Institution')
                    start_date = edu.get('start_date', '')
                    end_date = edu.get('end_date', '')
                    
                    p = doc.add_paragraph()
                    run = p.add_run(degree)
                    run.bold = True
                    run.font.size = Pt(11)
                    
                    p = doc.add_paragraph()
                    date_str = f" | {start_date} - {end_date}" if start_date else ""
                    p.add_run(f"{institution}{date_str}").italic = True
                    
                    if edu.get('description'):
                        doc.add_paragraph(edu['description'])
                    
                    doc.add_paragraph()  # Spacing
            
            # Skills
            if skills:
                doc.add_heading('Skills', level=2)
                if isinstance(skills, list):
                    skills_text = ', '.join(str(skill) for skill in skills)
                    doc.add_paragraph(skills_text)
                else:
                    doc.add_paragraph(str(skills))
            
            # Projects
            if projects:
                doc.add_heading('Projects', level=2)
                for project in projects:
                    name = project.get('name', 'Project Name')
                    
                    p = doc.add_paragraph()
                    run = p.add_run(name)
                    run.bold = True
                    run.font.size = Pt(11)
                    
                    if project.get('description'):
                        doc.add_paragraph(project['description'])
                    
                    if project.get('url'):
                        p = doc.add_paragraph()
                        p.add_run(f"Link: {project['url']}")
                    
                    doc.add_paragraph()  # Spacing
            
            # Save to bytes
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                doc.save(temp_docx.name)
                temp_docx_path = temp_docx.name
            
            try:
                with open(temp_docx_path, 'rb') as f:
                    docx_bytes = f.read()
            finally:
                try:
                    os.unlink(temp_docx_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temporary DOCX file: {e}")
            
            filename = f"cv-{cv.id}-v{cv.version}.docx"
            
            logger.info(f"Successfully generated DOCX for CV {cv.id}")
            return docx_bytes, filename
            
        except Exception as e:
            logger.error(f"Error generating DOCX for CV {cv.id}: {e}", exc_info=True)
            raise
