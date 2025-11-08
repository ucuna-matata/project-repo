"""Service for exporting CV to PDF and DOCX formats."""
from io import BytesIO
from datetime import datetime
from typing import Any, Tuple

from django.template.loader import render_to_string
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CVExportService:
    """Service for exporting CV to various formats."""

    @staticmethod
    def export_to_pdf(cv: Any) -> Tuple[BytesIO, str]:
        """
        Export CV to PDF format.

        Args:
            cv: CV model instance

        Returns:
            Tuple of (BytesIO buffer, filename)
        """
        # Render HTML template
        html_content = CVExportService._render_cv_html(cv)

        # Convert HTML to PDF using WeasyPrint
        pdf_buffer = BytesIO()
        HTML(string=html_content).write_pdf(pdf_buffer)
        pdf_buffer.seek(0)

        # Generate filename
        filename = CVExportService._generate_filename(cv, 'pdf')

        return pdf_buffer, filename

    @staticmethod
    def export_to_docx(cv: Any) -> Tuple[BytesIO, str]:
        """
        Export CV to DOCX format.

        Args:
            cv: CV model instance

        Returns:
            Tuple of (BytesIO buffer, filename)
        """
        document = Document()
        sections = cv.sections or {}

        # Set up document styles
        style = document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Personal Information
        personal = sections.get('personal', {})
        if personal:
            # Name as title
            if personal.get('name'):
                heading = document.add_heading(personal['name'], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contact info
            contact_parts = []
            if personal.get('email'):
                contact_parts.append(personal['email'])
            if personal.get('phone'):
                contact_parts.append(personal['phone'])
            if personal.get('location'):
                contact_parts.append(personal['location'])

            if contact_parts:
                contact = document.add_paragraph(' | '.join(contact_parts))
                contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Links
            links = personal.get('links', {})
            if links:
                links_parts = []
                if links.get('linkedin'):
                    links_parts.append(f"LinkedIn: {links['linkedin']}")
                if links.get('github'):
                    links_parts.append(f"GitHub: {links['github']}")
                if links.get('portfolio'):
                    links_parts.append(f"Portfolio: {links['portfolio']}")

                if links_parts:
                    links_para = document.add_paragraph(' | '.join(links_parts))
                    links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_paragraph()  # Spacing

        # Summary/Objective
        summary = sections.get('summary', '')
        if summary:
            document.add_heading('Professional Summary', level=2)
            document.add_paragraph(summary)
            document.add_paragraph()  # Spacing

        # Experience
        experience = sections.get('experience', [])
        if experience:
            document.add_heading('Work Experience', level=2)
            for exp in experience:
                # Job title and company
                title_para = document.add_paragraph()
                title_run = title_para.add_run(exp.get('position') or exp.get('title', 'Position'))
                title_run.bold = True
                title_run.font.size = Pt(12)

                company_para = document.add_paragraph()
                company_run = company_para.add_run(exp.get('company', 'Company'))
                company_run.italic = True

                # Dates
                date_parts = []
                start = exp.get('start_date') or exp.get('start')
                end = exp.get('end_date') or exp.get('end')
                if start:
                    date_parts.append(start)
                if end:
                    date_parts.append(end)
                elif start:
                    date_parts.append('Present')

                if date_parts:
                    date_para = document.add_paragraph(' - '.join(date_parts))
                    date_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)

                # Description
                if exp.get('description'):
                    document.add_paragraph(exp['description'])

                document.add_paragraph()  # Spacing between entries

        # Education
        education = sections.get('education', [])
        if education:
            document.add_heading('Education', level=2)
            for edu in education:
                # Degree
                degree_para = document.add_paragraph()
                degree_run = degree_para.add_run(edu.get('degree', 'Degree'))
                degree_run.bold = True
                degree_run.font.size = Pt(12)

                # Institution
                if edu.get('institution'):
                    inst_para = document.add_paragraph()
                    inst_run = inst_para.add_run(edu['institution'])
                    inst_run.italic = True

                # Year/Dates
                if edu.get('start_date') and edu.get('end_date'):
                    year_para = document.add_paragraph(f"{edu['start_date']} - {edu['end_date']}")
                    year_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
                elif edu.get('year'):
                    year_para = document.add_paragraph(edu['year'])
                    year_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)

                # Description
                if edu.get('description'):
                    document.add_paragraph(edu['description'])

                document.add_paragraph()  # Spacing

        # Skills
        skills = sections.get('skills', [])
        if skills:
            document.add_heading('Skills', level=2)

            # Handle both formats: array of strings or array of objects
            if skills and isinstance(skills[0], str):
                # Simple list of skills as strings
                skills_para = document.add_paragraph(', '.join(skills))
            else:
                # Group skills by category
                skills_by_category = {}
                for skill in skills:
                    category = skill.get('category', 'General')
                    if category not in skills_by_category:
                        skills_by_category[category] = []
                    skills_by_category[category].append(skill.get('name', ''))

                for category, skill_names in skills_by_category.items():
                    cat_para = document.add_paragraph()
                    cat_run = cat_para.add_run(f"{category}: ")
                    cat_run.bold = True
                    cat_para.add_run(', '.join(skill_names))

            document.add_paragraph()  # Spacing

        # Projects
        projects = sections.get('projects', [])
        if projects:
            document.add_heading('Projects', level=2)
            for project in projects:
                # Project title
                title_para = document.add_paragraph()
                title_run = title_para.add_run(project.get('name') or project.get('title', 'Project'))
                title_run.bold = True
                title_run.font.size = Pt(12)

                # Description
                if project.get('description'):
                    document.add_paragraph(project['description'])

                # Technologies
                if project.get('technologies'):
                    tech_para = document.add_paragraph()
                    tech_para.add_run('Technologies: ').italic = True
                    if isinstance(project['technologies'], list):
                        tech_para.add_run(', '.join(project['technologies']))
                    else:
                        tech_para.add_run(project['technologies'])
                elif project.get('tech'):
                    tech_para = document.add_paragraph()
                    tech_para.add_run('Technologies: ').italic = True
                    tech_para.add_run(project['tech'])

                # Link
                if project.get('url'):
                    link_para = document.add_paragraph()
                    link_para.add_run('Link: ').italic = True
                    link_para.add_run(project['url'])
                elif project.get('link'):
                    link_para = document.add_paragraph()
                    link_para.add_run('Link: ').italic = True
                    link_para.add_run(project['link'])

                document.add_paragraph()  # Spacing

        # Save to buffer
        docx_buffer = BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)

        # Generate filename
        filename = CVExportService._generate_filename(cv, 'docx')

        return docx_buffer, filename

    @staticmethod
    def _render_cv_html(cv: Any) -> str:
        """Render CV as HTML using template."""
        # Normalize sections data
        sections = cv.sections or {}

        # Normalize skills format - convert array of strings to array of objects if needed
        if 'skills' in sections and sections['skills']:
            skills = sections['skills']
            if skills and isinstance(skills[0], str):
                # Convert array of strings to array of objects
                sections['skills'] = [{'name': skill, 'category': 'General'} for skill in skills]

        context = {
            'cv': cv,
            'sections': sections,
            'template_key': cv.template_key,
        }

        # Choose template based on template_key
        template_name = f'cv/{cv.template_key}.html'
        try:
            html = render_to_string(template_name, context)
        except:
            # Fallback to clean template
            html = render_to_string('cv/clean.html', context)

        return html

    @staticmethod
    def _generate_filename(cv: Any, extension: str) -> str:
        """Generate filename for exported CV."""
        # Sanitize title for filename
        title = cv.title.replace(' ', '_')
        title = ''.join(c for c in title if c.isalnum() or c in ('_', '-'))

        # Add timestamp
        timestamp = datetime.now().strftime('%Y%m%d')

        return f"{title}_{timestamp}.{extension}"

